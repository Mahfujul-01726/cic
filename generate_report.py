# -*- coding: utf-8 -*-
"""
Professional Word Report Generator — Chemical Image Classification
With ALL available figures embedded automatically.
"""

import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ────────────────────────────────────────────────────────────────
BASE        = r'd:\mahfuj\cic1\chemical_classification_models'
ACC_DIR     = os.path.join(BASE, 'Accuracy_Loss_Curves')
CM_DIR      = os.path.join(BASE, 'Confusion_Matrix')
ROC_DIR     = os.path.join(BASE, 'Results', 'ROC_Curves')
VIS_DIR     = os.path.join(BASE, 'Results', 'Visualizations')
SAMPLE_DIR  = os.path.join(BASE, 'samples and distrubutation')
RAW_DIR     = os.path.join(SAMPLE_DIR, 'raw dataset')
PRE_DIR     = os.path.join(SAMPLE_DIR, 'preprocess dataset')

MODEL_ORDER = ['VGG16','DenseNet121','ResNet50V2','InceptionV3',
               'MobileNetV2','Xception','EfficientNetB4','ConvNeXtTiny','CustomCNN']

def img_path(folder, filename):
    p = os.path.join(folder, filename)
    return p if os.path.exists(p) else None

# ── Document setup ────────────────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

# ── Style helpers ─────────────────────────────────────────────────────────
def set_font(run, name='Calibri', size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold      = bold
    run.italic    = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level=1, color=(0, 70, 127)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.name = 'Calibri'
    return p

def add_para(text, bold=False, italic=False, font_size=11,
             align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size=font_size, bold=bold, italic=italic, color=color)
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def add_code_block(code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p

def shade_row(row, fill_hex):
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_hex)
        tcPr.append(shd)

def set_cell_text(cell, text, bold=False, size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_hr():
    p = doc.add_paragraph()
    pPr   = p._p.get_or_add_pPr()
    pBdr  = OxmlElement('w:pBdr')
    bot   = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '4472C4')
    pBdr.append(bot)
    pPr.append(pBdr)

def insert_figure(img_file, caption, width=Inches(5.5)):
    """Insert an image (if it exists) with a centred caption."""
    if img_file and os.path.exists(img_file):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_file, width=width)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('[Image not found — screenshot placeholder]')
        set_font(run, italic=True, size=10, color=(180, 0, 0))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_font(r, italic=True, size=10, color=(89, 89, 89))
    doc.add_paragraph()   # spacing

# ═══════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_paragraph('\n\n')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LABORATORY REPORT')
set_font(run, name='Calibri', size=28, bold=True, color=(0, 70, 127))

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Chemical Image Classification')
set_font(run, name='Calibri', size=22, bold=True, color=(31, 78, 121))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('A Full Deep Learning Research Pipeline')
set_font(run, name='Calibri', size=14, italic=True, color=(89, 89, 89))

doc.add_paragraph()
add_hr()
doc.add_paragraph()

cover_tbl = doc.add_table(rows=5, cols=2)
cover_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cover_tbl.style = 'Table Grid'
cover_info = [
    ('Task',         'Chemical Image Classification using Deep Learning'),
    ('Platform',     'Google Colab (T4 GPU) + TensorFlow 2.20.0'),
    ('Dataset',      'Patent Chemical Images — 4 Classes, 16,000 Images'),
    ('Models Used',  'VGG16, DenseNet121, ResNet50V2, InceptionV3, MobileNetV2,\nXception, EfficientNetB4, ConvNeXtTiny, Custom CNN'),
    ('Date',         'July 2026'),
]
for i, (k, v) in enumerate(cover_info):
    row = cover_tbl.rows[i]
    shade_row(row, 'EBF3FB' if i % 2 == 0 else 'FFFFFF')
    set_cell_text(row.cells[0], k, bold=True, size=11, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(row.cells[1], v, size=11,   align=WD_ALIGN_PARAGRAPH.LEFT)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════
add_heading('Table of Contents', level=1)
toc_items = [
    ('1',    'Introduction & Task Requirements'),
    ('2',    'Dataset Overview & Directory Structure'),
    ('3',    'Sample Images from Every Class'),
    ('4',    'Class-wise Image Distribution (Bar Plot)'),
    ('5',    'Image Preprocessing & Augmentation'),
    ('6',    'Model Architecture & Design'),
    ('  6.1','Transfer Learning Framework'),
    ('  6.2','Pre-trained Models'),
    ('  6.3','Custom CNN Architecture'),
    ('7',    'Training Configuration & Callbacks'),
    ('8',    'Model Summaries'),
    ('9',    'Accuracy & Loss Curves'),
    ('10',   'Model Performance Metrics'),
    ('11',   'Confusion Matrix & Classification Report'),
    ('12',   'ROC Curve & AUC'),
    ('13',   'Grad-CAM Visualisation'),
    ('14',   'LIME Explainability'),
    ('15',   'Conclusion & Key Findings'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'  {num}   {title}')
    set_font(run, size=11)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 1: INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════
add_heading('1. Introduction & Task Requirements', level=1)
add_para(
    'This report documents a complete deep learning research pipeline for classifying chemical images '
    'extracted from scientific publications and patent documents. The dataset contains four distinct '
    'categories of chemical imagery, and the objective was to design, train, and evaluate multiple deep '
    'learning architectures — both pre-trained transfer learning models and a custom-built CNN.', font_size=11
)
doc.add_paragraph()
add_para('The following tasks were assigned and performed:', bold=True, font_size=11)
tasks = [
    'Show sample images from every class.',
    'Show class-wise image quantity across train/validation/test sets using a bar plot.',
    'Perform image preprocessing and data augmentation.',
    'Use pre-trained models: VGG16, DenseNet121, ResNet50V2, InceptionV3, MobileNetV2, '
    'Xception, EfficientNetB4, and ConvNeXtTiny — with architectural modifications.',
    'Build and train a Custom CNN; ensemble models if required.',
    'Show model summaries.',
    'Plot accuracy and loss curves for training and validation sets.',
    'Compute Accuracy, Precision, Recall, F1-Score, Confusion Matrix, AUC, and ROC Curves.',
    'Generate Grad-CAM heatmaps to visualise model attention regions.',
    'Generate LIME explanations for model predictions.',
]
for t in tasks:
    add_bullet(t)
add_hr()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 2: DATASET
# ═══════════════════════════════════════════════════════════════════════════
add_heading('2. Dataset Overview & Directory Structure', level=1)
add_para(
    'The dataset consists of chemical images sourced from patent documents and scientific journals. '
    'Images were standardised to PNG format before training. The dataset is split into three sets: '
    'Training, Validation, and Test — with perfectly balanced class distribution.', font_size=11
)
doc.add_paragraph()
add_para('Class Descriptions:', bold=True, font_size=11)
cls_tbl = doc.add_table(rows=1, cols=2)
cls_tbl.style = 'Table Grid'
shade_row(cls_tbl.rows[0], '1F497D')
set_cell_text(cls_tbl.rows[0].cells[0], 'Class Name',   bold=True, size=11, color=(255,255,255))
set_cell_text(cls_tbl.rows[0].cells[1], 'Description',  bold=True, size=11, color=(255,255,255))
for cls, desc in [
    ('one_molecule',        'Images containing a single molecular structure or compound diagram.'),
    ('reactions',           'Images depicting chemical reactions with reagents, arrows, and products.'),
    ('rest',                'Other chemical-related images not fitting the above categories.'),
    ('several_molecules',   'Images containing multiple molecular structures arranged together.'),
]:
    r = cls_tbl.add_row()
    set_cell_text(r.cells[0], cls,  bold=True, size=11, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(r.cells[1], desc, size=11,   align=WD_ALIGN_PARAGRAPH.LEFT)

doc.add_paragraph()
add_para('Dataset Split Summary:', bold=True, font_size=11)
cnt_tbl = doc.add_table(rows=1, cols=6)
cnt_tbl.style = 'Table Grid'
shade_row(cnt_tbl.rows[0], '2E75B6')
for i, h in enumerate(['Split','one_molecule','reactions','rest','several_molecules','Total']):
    set_cell_text(cnt_tbl.rows[0].cells[i], h, bold=True, size=10, color=(255,255,255))
for ri, rd in enumerate([
    ['Train',       '3,200','3,200','3,200','3,200','12,800'],
    ['Validation',  '400',  '400',  '400',  '400',  '1,600'],
    ['Test',        '400',  '400',  '400',  '400',  '1,600'],
    ['Grand Total', '4,000','4,000','4,000','4,000','16,000'],
]):
    row = cnt_tbl.add_row()
    shade_row(row, 'D6E4F0' if ri==3 else ('EBF3FB' if ri%2==0 else 'FFFFFF'))
    for ci, val in enumerate(rd):
        set_cell_text(row.cells[ci], val, bold=(ri==3), size=10)

doc.add_paragraph()
add_para('Directory Structure:', bold=True, font_size=11)
add_code_block(
    "dataset_png/\n"
    "  ├── train/\n"
    "  │     ├── one_molecule/      (3,200 images)\n"
    "  │     ├── reactions/         (3,200 images)\n"
    "  │     ├── rest/              (3,200 images)\n"
    "  │     └── several_molecules/ (3,200 images)\n"
    "  ├── validation/\n"
    "  │     └── [400 images per class]\n"
    "  └── test/\n"
    "        └── [400 images per class]"
)
add_hr()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 3: SAMPLE IMAGES
# ═══════════════════════════════════════════════════════════════════════════
add_heading('3. Sample Images from Every Class', level=1)
add_para(
    'Random sample images were displayed from each of the four classes across all three dataset '
    'splits. A grid of 4 samples per class was generated for each split using matplotlib.', font_size=11
)
doc.add_paragraph()
add_code_block(
    "SAMPLES_PER_CLASS = 4\n"
    "fig, axes = plt.subplots(NUM_CLASSES, SAMPLES_PER_CLASS, figsize=(16, 16))\n"
    "fig.suptitle('Training Set Samples', fontsize=22, fontweight='bold')\n"
    "for row, cls in enumerate(CLASSES):\n"
    "    cls_dir = os.path.join(TRAIN_DIR, cls)\n"
    "    imgs    = get_images(cls_dir)\n"
    "    chosen  = random.sample(imgs, min(SAMPLES_PER_CLASS, len(imgs)))\n"
    "    for col, path in enumerate(chosen):\n"
    "        ax.imshow(load_image_rgb(path))\n"
    "plt.tight_layout() ; plt.show()"
)
doc.add_paragraph()
add_heading('3.1  Raw Training Set Samples', level=2)
add_para('4 randomly selected images from each class in the training split:', font_size=11)
insert_figure(img_path(RAW_DIR, 'train_samples.png'),
              'Figure 1: Raw Training Set — 4 samples × 4 classes', width=Inches(5.8))

add_heading('3.2  Raw Validation Set Samples', level=2)
add_para('4 randomly selected images from each class in the validation split:', font_size=11)
insert_figure(img_path(RAW_DIR, 'validation_samples.png'),
              'Figure 2: Raw Validation Set — 4 samples × 4 classes', width=Inches(5.8))

add_heading('3.3  Raw Test Set Samples', level=2)
add_para('4 randomly selected images from each class in the test split:', font_size=11)
insert_figure(img_path(RAW_DIR, 'test_samples.png'),
              'Figure 3: Raw Test Set — 4 samples × 4 classes', width=Inches(5.8))
add_hr()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 4: BAR PLOT
# ═══════════════════════════════════════════════════════════════════════════
add_heading('4. Class-wise Image Distribution (Bar Plot)', level=1)
add_para(
    'A grouped bar chart was produced showing the number of images per class across all three splits. '
    'The dataset is perfectly balanced — every class has equal counts in each split, ensuring '
    'unbiased model training.', font_size=11
)
doc.add_paragraph()
add_code_block(
    "x = np.arange(len(CLASSES))  ;  width = 0.25\n"
    "splits = ['train','validation','test']\n"
    "colors = ['#2196F3','#4CAF50','#FF9800']\n"
    "for idx, (split, color) in enumerate(zip(splits, colors)):\n"
    "    values = [counts[split][cls] for cls in CLASSES]\n"
    "    ax.bar(x + idx*width, values, width, label=split.capitalize(), color=color)\n"
    "ax.set_title('Class-wise Image Distribution Across Splits')\n"
    "plt.savefig('step2_class_distribution.png', dpi=150)"
)
doc.add_paragraph()
add_para('Result — Class Distribution Bar Chart (Raw Dataset):', bold=True, font_size=11)
insert_figure(img_path(RAW_DIR, 'class_distribution.png'),
              'Figure 4: Class-wise Image Distribution Across Train / Validation / Test Splits',
              width=Inches(5.5))
add_hr()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 5: PREPROCESSING & AUGMENTATION
# ═══════════════════════════════════════════════════════════════════════════
add_heading('5. Image Preprocessing & Augmentation', level=1)
add_para(
    'All images were standardised and augmented using Keras ImageDataGenerator. '
    'The following pipeline was applied to the training set:', font_size=11
)
doc.add_paragraph()
aug_tbl = doc.add_table(rows=1, cols=2)
aug_tbl.style = 'Table Grid'
shade_row(aug_tbl.rows[0], '2E75B6')
set_cell_text(aug_tbl.rows[0].cells[0], 'Augmentation Parameter', bold=True, size=10, color=(255,255,255))
set_cell_text(aug_tbl.rows[0].cells[1], 'Value / Setting',        bold=True, size=10, color=(255,255,255))
for ri, (k, v) in enumerate([
    ('Rescaling',          '1/255 — normalise pixel values to [0, 1]'),
    ('Rotation Range',     '±20°'),
    ('Width Shift Range',  '15% of image width'),
    ('Height Shift Range', '15% of image height'),
    ('Shear Range',        '10%'),
    ('Zoom Range',         '20%'),
    ('Horizontal Flip',    'Enabled'),
    ('Vertical Flip',      'Enabled'),
    ('Fill Mode',          'Nearest'),
    ('Validation Split',   '20% of training images'),
]):
    row = aug_tbl.add_row()
    shade_row(row, 'EBF3FB' if ri%2==0 else 'FFFFFF')
    set_cell_text(row.cells[0], k, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(row.cells[1], v, size=10,   align=WD_ALIGN_PARAGRAPH.LEFT)

doc.add_paragraph()
add_code_block(
    "train_datagen = ImageDataGenerator(\n"
    "    rescale=1./255, rotation_range=20,\n"
    "    width_shift_range=0.15, height_shift_range=0.15,\n"
    "    shear_range=0.1, zoom_range=0.2,\n"
    "    horizontal_flip=True, vertical_flip=True,\n"
    "    validation_split=0.2, fill_mode='nearest'\n"
    ")\n"
    "# Generators\n"
    "train_gen = ...flow_from_directory(..., subset='training')   # 10,240 images\n"
    "val_gen   = ...flow_from_directory(..., subset='validation') #  2,560 images\n"
    "test_gen  = val_test_datagen.flow_from_directory(...)        #  1,600 images"
)
doc.add_paragraph()
add_heading('5.1  Preprocessed Class Distribution', level=2)
add_para(
    'After preprocessing, the class balance is maintained across all splits. '
    'The bar chart below shows the image counts per class after augmentation-ready PNG conversion:',
    font_size=11
)
insert_figure(img_path(PRE_DIR, 'Figure_Preprocessed_Class_Distribution.png'),
              'Figure 5: Preprocessed Dataset — Class Distribution (Train / Val / Test)',
              width=Inches(5.5))

add_heading('5.2  Preprocessed Training Set Samples', level=2)
add_para('Sample images from the training set after preprocessing (resized, normalised):', font_size=11)
insert_figure(img_path(PRE_DIR, 'train_preprocessed_samples.png'),
              'Figure 6: Preprocessed Training Set Samples — 4 samples × 4 classes',
              width=Inches(5.8))

add_heading('5.3  Preprocessed Validation Set Samples', level=2)
add_para('Sample images from the validation set after preprocessing:', font_size=11)
insert_figure(img_path(PRE_DIR, 'validation_preprocessed_samples.png'),
              'Figure 7: Preprocessed Validation Set Samples — 4 samples × 4 classes',
              width=Inches(5.8))

add_heading('5.4  Preprocessed Test Set Samples', level=2)
add_para('Sample images from the test set after preprocessing:', font_size=11)
insert_figure(img_path(PRE_DIR, 'test_preprocessed_samples.png'),
              'Figure 8: Preprocessed Test Set Samples — 4 samples × 4 classes',
              width=Inches(5.8))
add_hr()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 6: MODEL ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════
add_heading('6. Model Architecture & Design', level=1)

add_heading('6.1 Transfer Learning Framework', level=2)
add_para(
    'A unified function build_transfer_model() applies the same classification head '
    'to every pre-trained base model:', font_size=11
)
for item in [
    'Pre-trained base (ImageNet weights, include_top=False)',
    'Global Average Pooling 2D',
    'Batch Normalization',
    'Dense(512, ReLU) + Dropout(0.5)',
    'Dense(256, ReLU) + Dropout(0.25)',
    'Dense(4, Softmax) — output layer',
]:
    add_bullet(item)

doc.add_paragraph()
add_code_block(
    "def build_transfer_model(base_model_fn, input_shape=(224,224,3),\n"
    "                         num_classes=4, dropout=0.5, fine_tune_from=None):\n"
    "    base = base_model_fn(weights='imagenet', include_top=False,\n"
    "                         input_shape=input_shape)\n"
    "    base.trainable = False\n"
    "    inputs = Input(shape=input_shape)\n"
    "    x = base(inputs, training=False)\n"
    "    x = GlobalAveragePooling2D()(x)\n"
    "    x = BatchNormalization()(x)\n"
    "    x = Dense(512, activation='relu')(x)\n"
    "    x = Dropout(0.5)(x)\n"
    "    x = Dense(256, activation='relu')(x)\n"
    "    x = Dropout(0.25)(x)\n"
    "    outputs = Dense(num_classes, activation='softmax')(x)\n"
    "    model = Model(inputs, outputs)\n"
    "    # Selective fine-tuning of last N layers\n"
    "    model.compile(optimizer=Adam(1e-4),\n"
    "                  loss='categorical_crossentropy', metrics=['accuracy'])\n"
    "    return model"
)

add_heading('6.2 Pre-trained Models Used', level=2)
mdl_tbl = doc.add_table(rows=1, cols=4)
mdl_tbl.style = 'Table Grid'
shade_row(mdl_tbl.rows[0], '1F497D')
for i, h in enumerate(['Model','Input Size','Epochs / Fine-tune','Key Characteristic']):
    set_cell_text(mdl_tbl.rows[0].cells[i], h, bold=True, size=10, color=(255,255,255))
for ri, rd in enumerate([
    ('VGG16',          '224×224', '5 / fine_tune_from=8',  'Oxford VGG 16-layer deep CNN'),
    ('DenseNet121',    '224×224', '5 / fine_tune_from=20', 'Dense skip connections between all layers'),
    ('ResNet50V2',     '224×224', '5 / fine_tune_from=20', 'Improved residual learning (V2)'),
    ('InceptionV3',    '299×299', '3 / fine_tune_from=20', 'Multi-scale inception modules'),
    ('MobileNetV2',    '224×224', '3 / fine_tune_from=30', 'Lightweight depthwise separable convs'),
    ('Xception',       '299×299', '5 / fine_tune_from=20', 'Extreme inception — full depthwise convs'),
    ('EfficientNetB4', '224×224', '5 / fine_tune_from=30', 'Compound width/depth/resolution scaling'),
    ('ConvNeXtTiny',   '224×224', '5 / fine_tune_from=20', 'Modern pure ConvNet, ViT-inspired design'),
]):
    row = mdl_tbl.add_row()
    shade_row(row, 'EBF3FB' if ri%2==0 else 'FFFFFF')
    for ci, val in enumerate(rd):
        set_cell_text(row.cells[ci], val, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)

add_heading('6.3 Custom CNN Architecture', level=2)
add_para('A deep custom CNN was built with 4 convolutional blocks + a classification head:', font_size=11)
add_code_block(
    "def build_custom_cnn(input_shape=(224,224,3), num_classes=4):\n"
    "    # Block 1:  Conv2D(64) × 2  + BN + MaxPool + Dropout(0.20)\n"
    "    # Block 2:  Conv2D(128) × 2 + BN + MaxPool + Dropout(0.25)\n"
    "    # Block 3:  Conv2D(256) × 2 + BN + MaxPool + Dropout(0.30)\n"
    "    # Block 4:  Conv2D(512) × 2 + BN + MaxPool + Dropout(0.35)\n"
    "    # Head:     GlobalAvgPool → Dense(512)+BN+Drop(0.5)\n"
    "    #           → Dense(256)+Drop(0.3) → Dense(4, softmax)\n"
    "    model.compile(optimizer=Adam(1e-3), loss='categorical_crossentropy')"
)
add_hr()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 7: TRAINING CONFIG
# ═══════════════════════════════════════════════════════════════════════════
add_heading('7. Training Configuration & Callbacks', level=1)
cfg_tbl = doc.add_table(rows=1, cols=2)
cfg_tbl.style = 'Table Grid'
shade_row(cfg_tbl.rows[0], '2E75B6')
set_cell_text(cfg_tbl.rows[0].cells[0], 'Parameter', bold=True, size=10, color=(255,255,255))
set_cell_text(cfg_tbl.rows[0].cells[1], 'Value',     bold=True, size=10, color=(255,255,255))
for ri, (k, v) in enumerate([
    ('Optimizer',          'Adam (lr = 1e-4)'),
    ('Loss Function',      'Categorical Cross-Entropy'),
    ('Batch Size',         '32'),
    ('Epochs per Model',   '3 – 5 (with Early Stopping)'),
    ('Image Size (most)',  '224×224 px'),
    ('Image Size (Inception/Xception)', '299×299 px'),
    ('TensorFlow Version', '2.20.0'),
    ('Hardware',           'Google Colab T4 GPU'),
    ('Random Seed',        '42'),
]):
    row = cfg_tbl.add_row()
    shade_row(row, 'EBF3FB' if ri%2==0 else 'FFFFFF')
    set_cell_text(row.cells[0], k, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(row.cells[1], v, size=10,   align=WD_ALIGN_PARAGRAPH.LEFT)

doc.add_paragraph()
add_para('Callbacks Used:', bold=True, font_size=11)
for c in [
    'EarlyStopping (patience=7, monitor=val_accuracy, restore_best_weights=True)',
    'ReduceLROnPlateau (factor=0.3, patience=3, min_lr=1e-7, monitor=val_loss)',
    'ModelCheckpoint (saves best model as .keras file to Google Drive)',
]:
    add_bullet(c)
add_hr()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 8: MODEL SUMMARIES
# ═══════════════════════════════════════════════════════════════════════════
add_heading('8. Model Summaries', level=1)
add_para(
    'Detailed model summaries were generated using model.summary() and saved as text files. '
    'The summaries show every layer, its output shape, and the total trainable and non-trainable '
    'parameter counts. The table below shows the approximate parameter counts for each model:', font_size=11
)
doc.add_paragraph()
sum_tbl = doc.add_table(rows=1, cols=3)
sum_tbl.style = 'Table Grid'
shade_row(sum_tbl.rows[0], '1F497D')
for i, h in enumerate(['Model', 'Base Params (approx.)', 'Notes']):
    set_cell_text(sum_tbl.rows[0].cells[i], h, bold=True, size=10, color=(255,255,255))
for ri, rd in enumerate([
    ('VGG16',          '138 M',  'Heaviest model, no skip connections'),
    ('DenseNet121',    '7 M',    'Parameter-efficient due to dense connections'),
    ('ResNet50V2',     '25 M',   'Residual blocks reduce vanishing gradients'),
    ('InceptionV3',    '23 M',   'Multi-scale feature extraction'),
    ('MobileNetV2',    '3.4 M',  'Lightest model, designed for mobile'),
    ('Xception',       '22 M',   'Depthwise separable convolutions'),
    ('EfficientNetB4', '19 M',   'Compound scaling'),
    ('ConvNeXtTiny',   '28 M',   'Latest architecture, ViT-inspired'),
    ('Custom CNN',     '~10 M',  '4-block from-scratch CNN'),
]):
    row = sum_tbl.add_row()
    shade_row(row, 'EBF3FB' if ri%2==0 else 'FFFFFF')
    for ci, val in enumerate(rd):
        set_cell_text(row.cells[ci], val, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
add_hr()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 9: ACCURACY & LOSS CURVES  ← REAL FIGURES
# ═══════════════════════════════════════════════════════════════════════════
add_heading('9. Accuracy & Loss Curves', level=1)
add_para(
    'For each trained model, combined accuracy and loss curves are plotted showing both '
    'training and validation performance across epochs. These curves help diagnose overfitting '
    'or underfitting behaviour.', font_size=11
)
doc.add_paragraph()
add_code_block(
    "for name, history in histories.items():\n"
    "    plt.subplot(1,2,1)\n"
    "    plt.plot(history['accuracy'],     label='Train')\n"
    "    plt.plot(history['val_accuracy'], label='Validation')\n"
    "    plt.subplot(1,2,2)\n"
    "    plt.plot(history['loss'],     label='Train')\n"
    "    plt.plot(history['val_loss'], label='Validation')"
)
doc.add_paragraph()

acc_order = ['VGG16','DenseNet121','ResNet50V2','InceptionV3',
             'MobileNetV2','Xception','EfficientNetB4','ConvNeXtTiny','CustomCNN']
acc_labels = {
    'VGG16':         'Figure 1: VGG16 — Training & Validation Accuracy/Loss',
    'DenseNet121':   'Figure 2: DenseNet121 — Training & Validation Accuracy/Loss',
    'ResNet50V2':    'Figure 3: ResNet50V2 — Training & Validation Accuracy/Loss',
    'InceptionV3':   'Figure 4: InceptionV3 — Training & Validation Accuracy/Loss',
    'MobileNetV2':   'Figure 5: MobileNetV2 — Training & Validation Accuracy/Loss',
    'Xception':      'Figure 6: Xception — Training & Validation Accuracy/Loss',
    'EfficientNetB4':'Figure 7: EfficientNetB4 — Training & Validation Accuracy/Loss',
    'ConvNeXtTiny':  'Figure 8: ConvNeXtTiny — Training & Validation Accuracy/Loss',
    'CustomCNN':     'Figure 9: Custom CNN — Training & Validation Accuracy/Loss',
}
for m in acc_order:
    ip = img_path(ACC_DIR, f'{m}.png')
    insert_figure(ip, acc_labels.get(m, m), width=Inches(5.8))

add_hr()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 10: MODEL METRICS
# ═══════════════════════════════════════════════════════════════════════════
add_heading('10. Model Performance Metrics', level=1)
add_para(
    'All 9 models were evaluated on the held-out test set (1,600 images). '
    'Metrics computed: Accuracy, Precision (weighted), Recall (weighted), F1-Score (weighted).', font_size=11
)
doc.add_paragraph()
perf_tbl = doc.add_table(rows=1, cols=5)
perf_tbl.style = 'Table Grid'
shade_row(perf_tbl.rows[0], '1F497D')
for i, h in enumerate(['Model','Accuracy','Precision','Recall','F1-Score']):
    set_cell_text(perf_tbl.rows[0].cells[i], h, bold=True, size=10, color=(255,255,255))
metrics_data = [
    ('VGG16',          0.9106, 0.9115, 0.9106, 0.9104),
    ('ResNet50V2',     0.8988, 0.8990, 0.8988, 0.8985),
    ('DenseNet121',    0.8913, 0.8912, 0.8913, 0.8910),
    ('InceptionV3',    0.8863, 0.8901, 0.8863, 0.8866),
    ('Xception',       0.8088, 0.8093, 0.8088, 0.8089),
    ('ConvNeXtTiny',   0.7688, 0.8152, 0.7688, 0.7763),
    ('MobileNetV2',    0.7688, 0.8338, 0.7688, 0.7552),
    ('EfficientNetB4', 0.3938, 0.5850, 0.3938, 0.3227),
    ('CustomCNN',      0.2694, 0.3267, 0.2694, 0.1504),
]
for ri, (m, a, p, r, f) in enumerate(metrics_data):
    row = perf_tbl.add_row()
    fill = 'E2EFDA' if ri==0 else ('EBF3FB' if ri%2==0 else 'FFFFFF')
    shade_row(row, fill)
    set_cell_text(row.cells[0], m, bold=(ri==0), size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    for ci, val in enumerate([a,p,r,f],1):
        set_cell_text(row.cells[ci], f'{val:.4f}', size=10)

doc.add_paragraph()
add_para(
    'Best Model: VGG16 achieved the highest accuracy (91.06%) and F1-Score (91.04%). '
    'EfficientNetB4 and CustomCNN underperformed — EfficientNet likely due to '
    'preprocessing incompatibility, and CustomCNN due to insufficient training data for '
    'a from-scratch model.',
    italic=True, font_size=11, color=(0, 100, 0)
)
add_hr()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 11: CONFUSION MATRIX  ← REAL FIGURES
# ═══════════════════════════════════════════════════════════════════════════
add_heading('11. Confusion Matrix & Classification Report', level=1)
add_para(
    'Seaborn heatmap confusion matrices were generated for all 9 models. Each cell shows '
    'the number of test images predicted as a given class for each true class.', font_size=11
)
doc.add_paragraph()
add_para('VGG16 Classification Report (Best Model):', bold=True, font_size=11)
add_code_block(
    "                   precision    recall  f1-score   support\n\n"
    "     one_molecule       0.89      0.94      0.92       400\n"
    "        reactions       0.94      0.86      0.90       400\n"
    "             rest       0.93      0.94      0.94       400\n"
    "several_molecules       0.89      0.89      0.89       400\n\n"
    "         accuracy                           0.91      1600\n"
    "        macro avg       0.91      0.91      0.91      1600\n"
    "     weighted avg       0.91      0.91      0.91      1600"
)
doc.add_paragraph()

cm_order  = ['VGG16','DenseNet121','ResNet50V2','InceptionV3',
             'MobileNetV2','Xception','EfficientNetB4','ConvNeXtTiny','CustomCNN']
cm_labels = {m: f'Figure {10+i}: {m} — Confusion Matrix'
             for i, m in enumerate(cm_order)}
for m in cm_order:
    ip = img_path(CM_DIR, f'{m}_Confusion_Matrix.png')
    insert_figure(ip, cm_labels[m], width=Inches(4.5))

add_hr()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 12: ROC & AUC  ← REAL FIGURES
# ═══════════════════════════════════════════════════════════════════════════
add_heading('12. ROC Curve & AUC', level=1)
add_para(
    'Multi-class ROC curves were computed using the One-vs-Rest (OvR) strategy with label '
    'binarization. The AUC (Area Under the Curve) was macro-averaged across all 4 classes.', font_size=11
)
doc.add_paragraph()
auc_tbl = doc.add_table(rows=1, cols=2)
auc_tbl.style = 'Table Grid'
shade_row(auc_tbl.rows[0], '2E75B6')
set_cell_text(auc_tbl.rows[0].cells[0], 'Model',          bold=True, size=10, color=(255,255,255))
set_cell_text(auc_tbl.rows[0].cells[1], 'AUC (Macro-Avg)',bold=True, size=10, color=(255,255,255))
for ri, (m, a) in enumerate([
    ('VGG16',0.9890),('DenseNet121',0.9832),('ResNet50V2',0.9821),
    ('InceptionV3',0.9846),('Xception',0.9554),('MobileNetV2',0.9614),
    ('ConvNeXtTiny',0.9459),('EfficientNetB4',0.7105),('CustomCNN',0.6743),
]):
    row = auc_tbl.add_row()
    shade_row(row, 'E2EFDA' if ri==0 else ('EBF3FB' if ri%2==0 else 'FFFFFF'))
    set_cell_text(row.cells[0], m,           bold=(ri==0), size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(row.cells[1], f'{a:.4f}',  size=10)

doc.add_paragraph()
roc_order  = ['VGG16','DenseNet121','ResNet50V2','InceptionV3',
              'MobileNetV2','Xception','EfficientNetB4','ConvNeXtTiny','CustomCNN']
roc_labels = {m: f'Figure {19+i}: {m} — ROC Curve (OvR)'
              for i, m in enumerate(roc_order)}
for m in roc_order:
    ip = img_path(ROC_DIR, f'{m}_ROC.png')
    insert_figure(ip, roc_labels[m], width=Inches(5.0))

add_hr()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 13: GRAD-CAM  ← REAL FIGURE
# ═══════════════════════════════════════════════════════════════════════════
add_heading('13. Grad-CAM Visualisation', level=1)
add_para(
    'Gradient-weighted Class Activation Mapping (Grad-CAM) was applied to VGG16 '
    '(using the block5_conv3 convolutional layer) to visualise which image regions '
    'the model focused on during classification.', font_size=11
)
doc.add_paragraph()
add_para('How Grad-CAM Works:', bold=True, font_size=11)
add_para(
    'Grad-CAM computes the gradient of the predicted class score with respect to the last '
    'convolutional layer\'s feature maps. These gradients are global-average-pooled to produce '
    'class-discriminative weights, which are combined with the feature maps to create a heatmap. '
    'The heatmap is overlaid on the original image using a jet colour map.', font_size=11
)
doc.add_paragraph()
add_code_block(
    "grad_model = tf.keras.models.Model(\n"
    "    inputs=[base_model.input],\n"
    "    outputs=[base_model.get_layer('block5_conv3').output, base_model.output]\n"
    ")\n"
    "with tf.GradientTape() as tape:\n"
    "    conv_outputs, predictions = grad_model(img_tensor)\n"
    "    class_channel = tf.gather(predictions[0], pred_index)\n"
    "grads       = tape.gradient(class_channel, conv_outputs)\n"
    "pooled_grads = tf.reduce_mean(grads, axis=(0,1,2))\n"
    "heatmap     = conv_outputs[0] @ pooled_grads[..., tf.newaxis]\n"
    "heatmap     = tf.maximum(heatmap, 0) / tf.math.reduce_max(heatmap)"
)
doc.add_paragraph()
gradcam_ip = img_path(VIS_DIR, 'VGG16_GradCAM.png')
insert_figure(gradcam_ip, 'Figure 28: Grad-CAM Heatmaps — VGG16 (one row per class, original + overlay)', width=Inches(5.5))
add_hr()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 14: LIME  ← REAL FIGURE
# ═══════════════════════════════════════════════════════════════════════════
add_heading('14. LIME Explainability', level=1)
add_para(
    'LIME (Local Interpretable Model-agnostic Explanations) was applied to the VGG16 model. '
    'LIME perturbs the input image into superpixel segments and fits a local linear model '
    'around each prediction to identify which image regions support or contradict the prediction.', font_size=11
)
doc.add_paragraph()
add_para('Output per class (3 subplot columns):', bold=True, font_size=11)
for item in [
    'Original Image — raw test image from the class.',
    'Positive Regions (Green) — superpixels that supported the top predicted class.',
    'Negative Regions (Red)   — superpixels that worked against the top prediction.',
]:
    add_bullet(item)
doc.add_paragraph()
add_code_block(
    "explainer = lime_image.LimeImageExplainer()\n"
    "explanation = explainer.explain_instance(\n"
    "    img_array.astype('double'),\n"
    "    predict_fn,      # model.predict wrapper\n"
    "    top_labels=4, hide_color=0, num_samples=1000\n"
    ")\n"
    "temp, mask = explanation.get_image_and_mask(\n"
    "    top_label, positive_only=True, hide_rest=False)"
)
doc.add_paragraph()
lime_ip = img_path(VIS_DIR, 'Lime_VGG16_Fixed.png')
insert_figure(lime_ip, 'Figure 29: LIME Explanations — VGG16 (all 4 classes)', width=Inches(5.5))
add_hr()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 15: CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════
add_heading('15. Conclusion & Key Findings', level=1)
add_para(
    'This report presents a comprehensive deep learning pipeline for chemical image classification. '
    'Nine different architectures were trained and evaluated:', font_size=11
)
doc.add_paragraph()
for c in [
    'VGG16 achieved the best overall performance (Accuracy: 91.06%, AUC: 0.9890), making it the most '
     'suitable model for this task.',
    'ResNet50V2 and DenseNet121 also performed strongly (~89–90% accuracy), confirming that residual '
     'and dense connection architectures transfer well to chemical imagery.',
    'InceptionV3 (88.6%, AUC: 0.9846) and Xception (80.9%) showed competitive results using 299×299 input.',
    'MobileNetV2 and ConvNeXtTiny both achieved 76.9% accuracy; MobileNetV2 suits deployment where '
     'model size is constrained.',
    'EfficientNetB4 underperformed significantly (39.4%) — likely due to its built-in normalization '
     'conflicting with the rescale=1./255 generator pipeline.',
    'The Custom CNN achieved only 26.9% accuracy, showing that training a deep CNN from scratch on '
     '12,800 images is insufficient without pre-trained features.',
    'Grad-CAM confirmed VGG16 correctly focused on molecular bonds, reaction arrows, and structural '
     'patterns — validating genuine feature learning.',
    'LIME explanations showed clear positive (green) regions aligned with key chemical structures, '
     'providing interpretable evidence for each prediction.',
]:
    add_bullet(c)

doc.add_paragraph()
add_para('Final Model Ranking by Accuracy:', bold=True, font_size=11)
rank_tbl = doc.add_table(rows=1, cols=4)
rank_tbl.style = 'Table Grid'
shade_row(rank_tbl.rows[0], '1F497D')
for i, h in enumerate(['Rank','Model','Accuracy','AUC']):
    set_cell_text(rank_tbl.rows[0].cells[i], h, bold=True, size=10, color=(255,255,255))
for ri, (rk,m,a,au) in enumerate([
    ('1st','VGG16',         '91.06%','0.9890'),
    ('2nd','ResNet50V2',    '89.88%','0.9821'),
    ('3rd','DenseNet121',   '89.13%','0.9832'),
    ('4th','InceptionV3',   '88.63%','0.9846'),
    ('5th','Xception',      '80.88%','0.9554'),
    ('6th','MobileNetV2',   '76.88%','0.9614'),
    ('6th','ConvNeXtTiny',  '76.88%','0.9459'),
    ('8th','EfficientNetB4','39.38%','0.7105'),
    ('9th','Custom CNN',    '26.94%','0.6743'),
]):
    row = rank_tbl.add_row()
    shade_row(row, 'E2EFDA' if ri==0 else ('EBF3FB' if ri%2==0 else 'FFFFFF'))
    for ci, val in enumerate([rk,m,a,au]):
        set_cell_text(row.cells[ci], val, bold=(ri==0), size=10, align=WD_ALIGN_PARAGRAPH.LEFT)

doc.add_paragraph()
add_hr()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— End of Report —')
set_font(run, italic=True, size=12, color=(89,89,89))

# ── SAVE ──────────────────────────────────────────────────────────────────
out_path = r'd:\mahfuj\cic1\Chemical_Image_Classification_Report.docx'
doc.save(out_path)
print(f'✅  Report saved: {out_path}')
print(f'    Total pages (estimated): ~60+')
